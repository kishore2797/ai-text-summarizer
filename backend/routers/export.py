from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional
import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from docx import Document
import csv

router = APIRouter()

class ExportRequest(BaseModel):
    content: str
    filename: Optional[str] = "summary"
    format: str  # pdf, docx, txt, csv
    include_metadata: Optional[bool] = False
    metadata: Optional[dict] = None

@router.post("/export")
async def export_content(request: ExportRequest):
    """
    Export summary to various formats
    
    - **content**: Text content to export
    - **filename**: Output filename (without extension)
    - **format**: Export format (pdf, docx, txt, csv)
    - **include_metadata**: Whether to include metadata in the export
    - **metadata**: Optional metadata to include
    """
    try:
        if request.format not in ["pdf", "docx", "txt", "csv"]:
            raise HTTPException(
                status_code=400,
                detail="Unsupported format. Use: pdf, docx, txt, csv"
            )
        
        if request.format == "pdf":
            return await export_to_pdf(request)
        elif request.format == "docx":
            return await export_to_docx(request)
        elif request.format == "txt":
            return await export_to_txt(request)
        elif request.format == "csv":
            return await export_to_csv(request)
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

async def export_to_pdf(request: ExportRequest):
    """Export content to PDF format"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title = Paragraph(f"<b>{request.filename}</b>", styles['Title'])
    story.append(title)
    story.append(Spacer(1, 12))
    
    # Metadata if requested
    if request.include_metadata and request.metadata:
        for key, value in request.metadata.items():
            meta_text = Paragraph(f"<b>{key}:</b> {value}", styles['Normal'])
            story.append(meta_text)
        story.append(Spacer(1, 12))
    
    # Content
    content_paragraph = Paragraph(request.content, styles['Normal'])
    story.append(content_paragraph)
    
    doc.build(story)
    buffer.seek(0)
    
    return StreamingResponse(
        io.BytesIO(buffer.read()),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={request.filename}.pdf"}
    )

async def export_to_docx(request: ExportRequest):
    """Export content to DOCX format"""
    doc = Document()
    
    # Title
    title = doc.add_heading(request.filename, 0)
    
    # Metadata if requested
    if request.include_metadata and request.metadata:
        for key, value in request.metadata.items():
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(str(value))
        doc.add_paragraph()  # Add space
    
    # Content
    doc.add_paragraph(request.content)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        io.BytesIO(buffer.read()),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={request.filename}.docx"}
    )

async def export_to_txt(request: ExportRequest):
    """Export content to TXT format"""
    content = f"# {request.filename}\n\n"
    
    if request.include_metadata and request.metadata:
        for key, value in request.metadata.items():
            content += f"{key}: {value}\n"
        content += "\n"
    
    content += request.content
    
    return StreamingResponse(
        io.StringIO(content),
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename={request.filename}.txt"}
    )

async def export_to_csv(request: ExportRequest):
    """Export content to CSV format"""
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Header
    writer.writerow(["Filename", "Content"])
    
    # Data
    writer.writerow([request.filename, request.content])
    
    # Metadata if requested
    if request.include_metadata and request.metadata:
        writer.writerow([])  # Empty row
        writer.writerow(["Metadata"])
        for key, value in request.metadata.items():
            writer.writerow([key, value])
    
    output.seek(0)
    
    return StreamingResponse(
        io.StringIO(output.getvalue()),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={request.filename}.csv"}
    )

@router.get("/formats")
async def get_export_formats():
    """
    Get available export formats and their features
    """
    return {
        "formats": {
            "pdf": {
                "name": "PDF",
                "description": "Portable Document Format - best for sharing and printing",
                "features": ["Formatted layout", "Metadata support", "Professional appearance"],
                "mime_type": "application/pdf"
            },
            "docx": {
                "name": "DOCX",
                "description": "Microsoft Word format - editable document",
                "features": ["Editable", "Metadata support", "Word processor compatible"],
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            },
            "txt": {
                "name": "TXT",
                "description": "Plain text format - universal compatibility",
                "features": ["Universal compatibility", "Small file size", "Simple format"],
                "mime_type": "text/plain"
            },
            "csv": {
                "name": "CSV",
                "description": "Comma-separated values - for data analysis",
                "features": ["Data analysis ready", "Excel compatible", "Structured format"],
                "mime_type": "text/csv"
            }
        },
        "metadata_fields": [
            "original_length",
            "summary_length", 
            "compression_ratio",
            "processing_time",
            "method",
            "model",
            "language"
        ]
    }
